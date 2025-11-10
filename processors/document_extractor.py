# DEPENDENCIES
import io
import os
import re
import mimetypes
from typing import Any
from typing import Dict
from typing import List
from pathlib import Path
from typing import Tuple
from loguru import logger
from typing import Optional 
from dataclasses import dataclass


# Document processing libraries
try:
    # PyMuPDF - Primary PDF Extractor
    import fitz  
    PYPDF_AVAILABLE = True
    logger.info("PyMuPDF available for high-quality PDF extraction")

except ImportError:
    logger.warning("PyMuPDF not available. Install: pip install PyMuPDF")
    PYPDF_AVAILABLE = False

try:
    # Fallback 1
    import pdfplumber  
    PDFPLUMBER_AVAILABLE = True

except ImportError:
    logger.warning("pdfplumber not available. Install: pip install pdfplumber")
    PDFPLUMBER_AVAILABLE = False

try:
    # Fallback 2
    import PyPDF2  
    PYPDF2_AVAILABLE = True

except ImportError:
    logger.warning("PyPDF2 not available. Install: pip install PyPDF2")
    PYPDF2_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True

except ImportError:
    logger.warning("python-docx not available. Install: pip install python-docx")
    DOCX_AVAILABLE = False

try:
    import chardet
    CHARDET_AVAILABLE = True

except ImportError:
    logger.warning("chardet not available. Install: pip install chardet")
    CHARDET_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True

except ImportError:
    logger.warning("BeautifulSoup not available. Install: pip install beautifulsoup4")
    BS4_AVAILABLE = False


@dataclass
class ExtractedDocument:
    """
    Container for extracted document content with metadata
    """
    text              : str
    file_path         : Optional[str]
    file_type         : str
    file_size_bytes   : int
    page_count        : int
    extraction_method : str
    metadata          : Dict[str, Any]
    is_success        : bool
    error_message     : Optional[str]
    warnings          : List[str]
    

    def to_dict(self) -> Dict[str, Any]:
        """
        Convert to dictionary for JSON serialization
        """
        return {"text_length"        : len(self.text),
                "file_type"          : self.file_type,
                "file_size_bytes"    : self.file_size_bytes,
                "page_count"         : self.page_count,
                "extraction_method"  : self.extraction_method,
                "metadata"           : self.metadata,
                "is_success"         : self.is_success,
                "error_message"      : self.error_message,
                "warnings"           : self.warnings,
               }


class DocumentExtractor:
    """
    Extracts text from various document formats for AI detection processing

    Supported Formats:
    - Plain text (.txt, .md, .log)
    - PDF documents (.pdf) - Uses PyMuPDF as primary extractor
    - Microsoft Word (.doc, .docx)
    - Rich Text Format (.rtf)
    - HTML files (.html, .htm)

    Features:
    - Robust error handling
    - Encoding detection
    - Metadata extraction
    - Page/section preservation
    - Memory-efficient processing
    """
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {'.txt', '.text', '.md', '.markdown', '.log', '.csv', '.pdf', '.docx', '.doc',  '.rtf', '.html', '.htm'}
    
    # Text file extensions
    TEXT_EXTENSIONS      = {'.txt', '.text', '.md', '.markdown', '.log', '.csv'}
    
    # Maximum file size (50 MB default)
    MAX_FILE_SIZE        = 50 * 1024 * 1024

    
    def __init__(self, max_file_size: int = MAX_FILE_SIZE, extract_metadata: bool = True):
        """
        Initialize document extractor
        
        Arguments:
        ----------
            max_file_size      { int }  : Maximum file size in bytes

            extract_metadata   { bool } : Extract document metadata
        """
        self.max_file_size      = max_file_size
        self.extract_metadata   = extract_metadata
        
        logger.info(f"DocumentExtractor initialized (max_size={max_file_size/1024/1024:.1f}MB)")
    

    def extract(self, file_path: str) -> ExtractedDocument:
        """
        Extract text from document
        
        Arguments:
        ----------
            file_path     { str } : Path to the document file
            
        Returns:
        --------
            { ExtractedDocument } : ExtractedDocument object with extracted text and metadata
        """
        try:
            file_path         = Path(file_path)
            
            # Validate file
            validation_result = self._validate_file(file_path)

            if not validation_result[0]:
                return self._create_error_result(file_path = str(file_path),
                                                 error     = validation_result[1],
                                                )
            
            # Get file info
            file_size = file_path.stat().st_size
            file_ext  = file_path.suffix.lower()
            
            # Route to appropriate extractor
            if (file_ext in self.TEXT_EXTENSIONS):
                result = self._extract_text_file(file_path)
            
            elif (file_ext == '.pdf'):
                result = self._extract_pdf(file_path)
            
            elif (file_ext in {'.docx', '.doc'}):
                result = self._extract_word(file_path)
            
            elif (file_ext == '.rtf'):
                result = self._extract_rtf(file_path)
            
            elif (file_ext in {'.html', '.htm'}):
                result = self._extract_html(file_path)
            
            else:
                return self._create_error_result(file_path = str(file_path),
                                                 error     = f"Unsupported file type: {file_ext}",
                                                )
            
            # Add common metadata
            result.file_path       = str(file_path)
            result.file_size_bytes = file_size
            
            logger.info(f"Extracted {len(result.text)} chars from {file_path.name} using {result.extraction_method}")
            return result
            
        except Exception as e:
            logger.error(f"Error extracting document: {repr(e)}")
            return self._create_error_result(file_path = str(file_path) if file_path else None,
                                             error     = repr(e),
                                            )
    

    def extract_from_bytes(self, file_bytes: bytes, filename: str, mime_type: Optional[str] = None) -> ExtractedDocument:
        """
        Extract text from bytes (for file uploads)
        
        Arguments:
        ----------
            file_bytes { bytes } : File content as bytes

            filename    { str }  : Original filename
            
            mime_type   { str }  : MIME type (optional)
            
        Returns:
        --------
            { ExtractedDocument} : ExtractedDocument object
        """
        try:
            # Determine file type
            file_ext = Path(filename).suffix.lower()
            
            if file_ext not in self.SUPPORTED_EXTENSIONS:
                return self._create_error_result(file_path = filename,
                                                 error     = f"Unsupported file type: {file_ext}",
                                                )
            
            # Check size
            if (len(file_bytes) > self.max_file_size):
                return self._create_error_result(file_path = filename,
                                                 error     = f"File too large: {len(file_bytes)/1024/1024:.1f}MB"
                                                )
            
            # Route to appropriate extractor
            if (file_ext in self.TEXT_EXTENSIONS):
                result = self._extract_text_bytes(file_bytes, filename)
            
            elif (file_ext == '.pdf'):
                result = self._extract_pdf_bytes(file_bytes, filename)
            
            elif (file_ext in {'.docx', '.doc'}):
                result = self._extract_word_bytes(file_bytes, filename)
            
            elif (file_ext == '.rtf'):
                result = self._extract_rtf_bytes(file_bytes, filename)
            
            elif (file_ext in {'.html', '.htm'}):
                result = self._extract_html_bytes(file_bytes, filename)
            
            else:
                return self._create_error_result(file_path = filename,
                                                 error     = f"Unsupported file type: {file_ext}"
                                                )
            
            result.file_path       = filename
            result.file_size_bytes = len(file_bytes)
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting from bytes: {repr(e)}")
            return self._create_error_result(file_path = filename,
                                             error     = repr(e),
                                            )
    
    
    def _extract_text_file(self, file_path: Path) -> ExtractedDocument:
        """
        Extract text from plain text files
        """
        warnings = list()
        
        try:
            # Try to detect encoding
            encoding = 'utf-8'
            
            if CHARDET_AVAILABLE:
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
                    detected = chardet.detect(raw_data)
                    
                    if (detected['confidence'] > 0.7):
                        encoding = detected['encoding']
                        logger.debug(f"Detected encoding: {encoding} (confidence: {detected['confidence']})")
            
            # Read file with detected encoding
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()

            except UnicodeDecodeError:
                # Fallback to latin-1 (never fails)
                warnings.append(f"Failed to decode with {encoding}, using latin-1")
                with open(file_path, 'r', encoding = 'latin-1') as f:
                    text = f.read()
            
            return ExtractedDocument(text              = text,
                                     file_path         = str(file_path),
                                     file_type         = file_path.suffix,
                                     file_size_bytes   = file_path.stat().st_size,
                                     page_count        = 1,
                                     extraction_method = 'plain_text',
                                     metadata          = {'encoding': encoding},
                                     is_success        = True,
                                     error_message     = None,
                                     warnings          = warnings,
                                    )
            
        except Exception as e:
            return self._create_error_result(file_path = str(file_path), 
                                             error     = repr(e),
                                            )
    

    def _extract_text_bytes(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        """
        Extract text from bytes
        """
        warnings = list()
        
        try:
            # Detect encoding
            encoding = 'utf-8'
            
            if CHARDET_AVAILABLE:
                detected = chardet.detect(file_bytes)
                if (detected['confidence'] > 0.7):
                    encoding = detected['encoding']
            
            # Decode
            try:
                text = file_bytes.decode(encoding)

            except UnicodeDecodeError:
                warnings.append(f"Failed to decode with {encoding}, using latin-1")
                text = file_bytes.decode('latin-1')
            
            return ExtractedDocument(text              = text,
                                     file_path         = filename,
                                     file_type         = Path(filename).suffix,
                                     file_size_bytes   = len(file_bytes),
                                     page_count        = 1,
                                     extraction_method = 'plain_text',
                                     metadata          = {'encoding': encoding},
                                     is_success        = True,
                                     error_message     = None,
                                     warnings          = warnings,
                                    )
            
        except Exception as e:
            return self._create_error_result(file_path = filename, 
                                             error     = repr(e),
                                            )
    

    def _extract_pdf(self, file_path: Path) -> ExtractedDocument:
        """
        Extract text from PDF files in exactly this Priority order : PyMuPDF > pdfplumber > PyPDF2
        """
        if not any([PYPDF_AVAILABLE, PDFPLUMBER_AVAILABLE, PYPDF2_AVAILABLE]):
            return self._create_error_result(file_path = str(file_path),
                                             error     = "PDF libraries not installed. Install: pip install PyMuPDF",
                                            )
        
        warnings          = list()
        text              = ""
        page_count        = 0
        metadata          = dict()
        extraction_method = "unknown"
        
        # Try PyMuPDF first : best quality & performance
        if PYPDF_AVAILABLE:
            try:
                doc        = fitz.open(file_path)
                page_count = doc.page_count
                metadata   = doc.metadata
                
                for page_num in range(page_count):
                    page      = doc[page_num]
                    page_text = page.get_text()
                    
                    if page_text:
                        text += page_text + "\n\n"
                
                doc.close()
                extraction_method = "PyMuPDF"
                
                if text.strip():
                    logger.info(f"Successfully extracted PDF using PyMuPDF: {len(text)} chars")
                    return ExtractedDocument(text              = text.strip(),
                                             file_path         = str(file_path),
                                             file_type         = '.pdf',
                                             file_size_bytes   = file_path.stat().st_size,
                                             page_count        = page_count,
                                             extraction_method = extraction_method,
                                             metadata          = metadata,
                                             is_success        = True,
                                             error_message     = None,
                                             warnings          = warnings,
                                            )
                
                else:
                    warnings.append("PyMuPDF extracted no text")

            except Exception as e:
                warnings.append(f"PyMuPDF failed: {repr(e)}, trying pdfplumber")
        
        # Fallback-1: Try pdfplumber
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    metadata   = pdf.metadata or {}
                    
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        
                        if page_text:
                            text += page_text + "\n\n"
                
                extraction_method = "pdfplumber"
                
                if text.strip():
                    logger.info(f"Extracted PDF using pdfplumber: {len(text)} chars")
                    return ExtractedDocument(text              = text.strip(),
                                             file_path         = str(file_path),
                                             file_type         = '.pdf',
                                             file_size_bytes   = file_path.stat().st_size,
                                             page_count        = page_count,
                                             extraction_method = extraction_method,
                                             metadata          = metadata,
                                             is_success        = True,
                                             error_message     = None,
                                             warnings          = warnings,
                                            )
                
                else:
                    warnings.append("pdfplumber extracted no text")
            
            except Exception as e:
                warnings.append(f"pdfplumber failed: {repr(e)}, trying PyPDF2")
        
        # Fallback-2: Try PyPDF2
        if PYPDF2_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    reader     = PyPDF2.PdfReader(f)
                    page_count = len(reader.pages)
                    
                    if self.extract_metadata:
                        metadata = reader.metadata or {}
                    
                    for page in reader.pages:
                        page_text = page.extract_text()
                        
                        if page_text:
                            text += page_text + "\n\n"
                
                extraction_method = "PyPDF2"
                
                if not text.strip():
                    warnings.append("PDF appears to be image-based or encrypted")
                
                return ExtractedDocument(text              = text.strip(),
                                         file_path         = str(file_path),
                                         file_type         = '.pdf',
                                         file_size_bytes   = file_path.stat().st_size,
                                         page_count        = page_count,
                                         extraction_method = extraction_method,
                                         metadata          = metadata,
                                         is_success        = bool(text.strip()),
                                         error_message     = None if text.strip() else "No text extracted from PDF",
                                         warnings          = warnings,
                                        )
                
            except Exception as e:
                warnings.append(f"PyPDF2 failed: {repr(e)}")
        
        # All extractors failed
        return self._create_error_result(file_path = str(file_path), 
                                         error     = "All PDF extractors failed: " + "; ".join(warnings),
                                        )

    
    def _extract_pdf_bytes(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        """
        Extract text from PDF bytes in this priority order : PyMuPDF > pdfplumber > PyPDF2
        """
        if not any([PYPDF_AVAILABLE, PDFPLUMBER_AVAILABLE, PYPDF2_AVAILABLE]):
            return self._create_error_result(file_path = filename, 
                                             error     = "PDF libraries not installed",
                                            )
        
        warnings          = list()
        text              = ""
        page_count        = 0
        metadata          = dict()
        extraction_method = "unknown"
        
        try:
            # Primary: Try PyMuPDF first
            if PYPDF_AVAILABLE:
                try:
                    doc        = fitz.open(stream=file_bytes, filetype="pdf")
                    page_count = doc.page_count
                    metadata   = doc.metadata
                    
                    for page_num in range(page_count):
                        page      = doc[page_num]
                        page_text = page.get_text()
                        
                        if page_text:
                            text += page_text + "\n\n"
                    
                    doc.close()
                    extraction_method = "PyMuPDF"
                    
                    if text.strip():
                        return ExtractedDocument(text              = text.strip(),
                                                 file_path         = filename,
                                                 file_type         = '.pdf',
                                                 file_size_bytes   = len(file_bytes),
                                                 page_count        = page_count,
                                                 extraction_method = extraction_method,
                                                 metadata          = metadata,
                                                 is_success        = True,
                                                 error_message     = None,
                                                 warnings          = warnings,
                                                )
                    
                    else:
                        warnings.append("PyMuPDF extracted no text")

                except Exception as e:
                    warnings.append(f"PyMuPDF failed: {repr(e)}, trying pdfplumber")
            
            # Fallback-1: Try pdfplumber
            if PDFPLUMBER_AVAILABLE:
                try:
                    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                        page_count = len(pdf.pages)
                        metadata   = pdf.metadata or {}
                        
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            
                            if page_text:
                                text += page_text + "\n\n"
                    
                    extraction_method = "pdfplumber"
                    
                    if text.strip():
                        return ExtractedDocument(text              = text.strip(),
                                                 file_path         = filename,
                                                 file_type         = '.pdf',
                                                 file_size_bytes   = len(file_bytes),
                                                 page_count        = page_count,
                                                 extraction_method = extraction_method,
                                                 metadata          = metadata,
                                                 is_success        = True,
                                                 error_message     = None,
                                                 warnings          = warnings,
                                                )
                    
                    else:
                        warnings.append("pdfplumber extracted no text")

                except Exception as e:
                    warnings.append(f"pdfplumber failed: {repr(e)}, trying PyPDF2")
            
            # Fallback-2: Try PyPDF2
            if PYPDF2_AVAILABLE:
                reader     = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                page_count = len(reader.pages)
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    
                    if page_text:
                        text += page_text + "\n\n"
                
                extraction_method = "PyPDF2"
                
                return ExtractedDocument(text              = text.strip(),
                                         file_path         = filename,
                                         file_type         = '.pdf',
                                         file_size_bytes   = len(file_bytes),
                                         page_count        = page_count,
                                         extraction_method = extraction_method,
                                         metadata          = metadata,
                                         is_success        = bool(text.strip()),
                                         error_message     = None if text.strip() else "No text extracted",
                                         warnings          = warnings,
                                        )
            
        except Exception as e:
            return self._create_error_result(file_path = filename, 
                                             error     = repr(e),
                                            )
        
        # All extractors failed
        return self._create_error_result(file_path = filename, 
                                         error     = "All PDF extractors failed: " + "; ".join(warnings),
                                        )

    
    def _extract_word(self, file_path: Path) -> ExtractedDocument:
        """
        Extract text from Word documents
        """
        if not DOCX_AVAILABLE:
            return self._create_error_result(file_path = str(file_path),
                                             error     = "python-docx not installed",
                                            )
        
        try:
            doc        = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text       = "\n\n".join(paragraphs)
            
            # Extract metadata
            metadata   = dict()

            if self.extract_metadata:
                core_props = doc.core_properties
                metadata   = {'author'   : core_props.author,
                              'title'    : core_props.title,
                              'subject'  : core_props.subject,
                              'created'  : str(core_props.created) if core_props.created else None,
                              'modified' : str(core_props.modified) if core_props.modified else None,
                            }
            
            return ExtractedDocument(text              = text,
                                     file_path         = str(file_path),
                                     file_type         = file_path.suffix,
                                     file_size_bytes   = file_path.stat().st_size,
                                     page_count        = len(paragraphs),
                                     extraction_method = 'python-docx',
                                     metadata          = metadata,
                                     is_success        = True,
                                     error_message     = None,
                                     warnings          = [],
                                    )
            
        except Exception as e:
            return self._create_error_result(file_path = str(file_path), 
                                             error     = repr(e),
                                            )

    
    def _extract_word_bytes(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        """
        Extract text from Word document bytes
        """
        if not DOCX_AVAILABLE:
            return self._create_error_result(file_path = filename, 
                                             error     = "python-docx not installed",
                                            )
        
        try:
            doc        = DocxDocument(io.BytesIO(file_bytes))
            
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text       = "\n\n".join(paragraphs)
            
            metadata   = dict()

            if self.extract_metadata:
                core_props = doc.core_properties
                metadata   = {'author' : core_props.author,
                              'title'  : core_props.title,
                             }
            
            return ExtractedDocument(text              = text,
                                     file_path         = filename,
                                     file_type         = Path(filename).suffix,
                                     file_size_bytes   = len(file_bytes),
                                     page_count        = len(paragraphs),
                                     extraction_method = 'python-docx',
                                     metadata          = metadata,
                                     is_success        = True,
                                     error_message     = None,
                                     warnings          = [],
                                    )
            
        except Exception as e:
            return self._create_error_result(file_path = filename, 
                                             error     = repr(e),
                                            )

    
    def _extract_rtf(self, file_path: Path) -> ExtractedDocument:
        """
        Extract text from RTF files (basic implementation)
        """
        warnings = ["RTF extraction is basic, formatting may be lost"]
        
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()
            
            # Very basic RTF stripping (remove control words)
            text = re.sub(r'\\[a-z]+\d*\s?', '', content)
            text = re.sub(r'[{}]', '', text)
            text = text.strip()
            
            return ExtractedDocument(text              = text,
                                     file_path         = str(file_path),
                                     file_type         = '.rtf',
                                     file_size_bytes   = file_path.stat().st_size,
                                     page_count        = 1,
                                     extraction_method = 'basic_rtf',
                                     metadata          = {},
                                     is_success        = True,
                                     error_message     = None,
                                     warnings          = warnings,
                                    )
            
        except Exception as e:
            return self._create_error_result(file_path = str(file_path), 
                                             error     = repr(e),
                                            )

    
    def _extract_rtf_bytes(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        """
        Extract text from RTF bytes
        """
        warnings = ["RTF extraction is basic, formatting may be lost"]
        
        try:
            content = file_bytes.decode('latin-1')
            
            # Basic RTF stripping
            text    = re.sub(r'\\[a-z]+\d*\s?', '', content)
            text    = re.sub(r'[{}]', '', text)
            text    = text.strip()
            
            return ExtractedDocument(text              = text,
                                     file_path         = filename,
                                     file_type         = '.rtf',
                                     file_size_bytes   = len(file_bytes),
                                     page_count        = 1,
                                     extraction_method = 'basic_rtf',
                                     metadata          = {},
                                     is_success        = True,
                                     error_message     = None,
                                     warnings          = warnings,
                                    )
            
        except Exception as e:
            return self._create_error_result(file_path = filename, 
                                             error     = repr(e),
                                            )

    
    def _extract_html(self, file_path: Path) -> ExtractedDocument:
        """
        Extract text from HTML files
        """
        if not BS4_AVAILABLE:
            return self._create_error_result(file_path = str(file_path),
                                             error     = "BeautifulSoup not installed",
                                            )
        
        try:
            with open(file_path, 'r', encoding = 'utf-8', errors = 'ignore') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text  = soup.get_text(separator = '\n')
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            text  = '\n'.join(line for line in lines if line)
            
            return ExtractedDocument(text              = text,
                                     file_path         = str(file_path),
                                     file_type         = file_path.suffix,
                                     file_size_bytes   = file_path.stat().st_size,
                                     page_count        = 1,
                                     extraction_method = 'beautifulsoup',
                                     metadata          = {},
                                     is_success        = True,
                                     error_message     = None,
                                     warnings          = [],
                                    )
            
        except Exception as e:
            return self._create_error_result(file_path = str(file_path), 
                                             error     = repr(e),
                                            )

    
    def _extract_html_bytes(self, file_bytes: bytes, filename: str) -> ExtractedDocument:
        """
        Extract text from HTML bytes
        """
        if not BS4_AVAILABLE:
            return self._create_error_result(file_path = filename, 
                                             error     = "BeautifulSoup not installed",
                                            )
        
        try:
            content = file_bytes.decode('utf-8', errors = 'ignore')
            
            soup    = BeautifulSoup(content, 'html.parser')
            
            for script in soup(["script", "style"]):
                script.decompose()
            
            text  = soup.get_text(separator='\n')
            lines = (line.strip() for line in text.splitlines())
            text  = '\n'.join(line for line in lines if line)
            
            return ExtractedDocument(text              = text,
                                     file_path         = filename,
                                     file_type         = Path(filename).suffix,
                                     file_size_bytes   = len(file_bytes),
                                     page_count        = 1,
                                     extraction_method = 'beautifulsoup',
                                     metadata          = {},
                                     is_success        = True,
                                     error_message     = None,
                                     warnings          = [],
                                    )
            
        except Exception as e:
            return self._create_error_result(file_path = filename,
                                             error     = repr(e),
                                            )

    
    def _validate_file(self, file_path: Path) -> Tuple[bool, Optional[str]]:
        """
        Validate file before extraction
        """
        # Check if file exists
        if not file_path.exists():
            return False, f"File not found: {file_path}"
        
        # Check if it's a file
        if not file_path.is_file():
            return False, f"Not a file: {file_path}"
        
        # Check file size
        file_size = file_path.stat().st_size
        
        if (file_size > self.max_file_size):
            return False, f"File too large: {file_size/1024/1024:.1f}MB (max: {self.max_file_size/1024/1024:.1f}MB)"
        
        # Check file extension
        if (file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS):
            return False, f"Unsupported file type: {file_path.suffix}"
        
        return True, None
    

    def _create_error_result(self, file_path: Optional[str], error: str) -> ExtractedDocument:
        """
        Create error result
        """
        return ExtractedDocument(text              = "",
                                 file_path         = file_path,
                                 file_type         = Path(file_path).suffix if file_path else "unknown",
                                 file_size_bytes   = 0,
                                 page_count        = 0,
                                 extraction_method = "failed",
                                 metadata          = {},
                                 is_success        = False,
                                 error_message     = error,
                                 warnings          = [],
                                )


# Convenience Functions
def extract_text(file_path: str, **kwargs) -> ExtractedDocument:
    """
    Quick text extraction with default settings
    
    Arguments:
    ----------
        file_path { str }     : Path to document

        **kwargs              : Override settings
        
    Returns:
    --------
        { ExtractedDocument } : ExtractedDocument object
    """
    extractor = DocumentExtractor(**kwargs)

    return extractor.extract(file_path)


def extract_from_upload(file_bytes: bytes, filename: str, **kwargs) -> ExtractedDocument:
    """
    Extract text from uploaded file
    
    Arguments:
    ----------
        file_bytes { bytes }  : File content as bytes
        
        filename    { str }   : Original filename

        **kwargs              : Override settings
        
    Returns:
    --------
        { ExtractedDocument } : ExtractedDocument object
    """
    extractor = DocumentExtractor(**kwargs)

    return extractor.extract_from_bytes(file_bytes, filename)


# Export
__all__ = ['extract_text',
           'DocumentExtractor',
           'ExtractedDocument',
           'extract_from_upload',
          ]
